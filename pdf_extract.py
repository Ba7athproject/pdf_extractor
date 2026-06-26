import sys
import os

if sys.platform == 'win32':
    cuda_base = r"C:\Program Files\NVIDIA GPU Computing Toolkit\CUDA"
    if os.path.exists(cuda_base):
        for folder in os.listdir(cuda_base):
            for sub in ["bin", r"bin\x64"]:
                full_path = os.path.join(cuda_base, folder, sub)
                if os.path.isdir(full_path):
                    try:
                        os.add_dll_directory(full_path)
                    except Exception:
                        pass

from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QPushButton, QTextEdit, QLabel, QFileDialog,
    QLineEdit, QMessageBox, QSpinBox, QProgressBar
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QPalette, QColor, QFont
from pdf2image import convert_from_path
import pytesseract
import tempfile
import re
import clipboard
from docx import Document
import openpyxl

try:
    import cv2
    USE_OPENCV = True
    from PIL import Image
    import numpy as np
except ImportError:
    USE_OPENCV = False

def clean_docx_text(text):
    # Enlève les NULL et caractères de contrôle à l'exception de tab, LF, CR
    text = text.replace('\x00', '')
    text = re.sub(r'[\x01-\x08\x0B-\x0C\x0E-\x1F]', '', text)
    return text

def clean_xlsx_text(text):
    # Même traitement, certains caractères XLSX non supportés (ex: NULL)
    text = text.replace('\x00', '')
    text = re.sub(r'[\x01-\x08\x0B-\x0C\x0E-\x1F]', '', text)
    return text

def parse_pages(pages_str, total_pages):
    res = set()
    for part in pages_str.split(","):
        part = part.strip()
        if '-' in part:
            start, end = map(int, part.split("-"))
            for i in range(start, end + 1):
                if 1 <= i <= total_pages:
                    res.add(i)
        elif part.isdigit():
            i = int(part)
            if 1 <= i <= total_pages:
                res.add(i)
    return sorted(res)

class OCRWorker(QThread):
    progress = pyqtSignal(int, str)
    done = pyqtSignal(str, list)
    def __init__(self, pdf, indices, total_global_pages, global_start_idx, dpi=100, poppler_path=None):
        super().__init__()
        self.pdf = pdf
        self.indices = indices
        self.dpi = dpi
        self.poppler_path = poppler_path
        self._cancelled = False
        self.total_global_pages = total_global_pages
        self.global_start_idx = global_start_idx

    def cancel(self):
        self._cancelled = True

    def run(self):
        txt = ""
        tables = []
        tmp_dir = tempfile.mkdtemp()
        for i, page_no in enumerate(self.indices):
            if self._cancelled:
                return
            overall_idx = self.global_start_idx + i
            percent = int((overall_idx+1) / self.total_global_pages * 100)
            self.progress.emit(percent, f"🕶️ OCR page {page_no} ({overall_idx+1}/{self.total_global_pages})")
            imgs = convert_from_path(self.pdf, dpi=self.dpi, first_page=page_no, last_page=page_no, poppler_path=self.poppler_path)
            if not imgs:
                continue
            img = imgs[0]
            if USE_OPENCV:
                img_np = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
                gpu_img = cv2.cuda_GpuMat()
                gpu_img.upload(img_np)
                gpu_gray = cv2.cuda.cvtColor(gpu_img, cv2.COLOR_BGR2GRAY)
                # gpu_gray = cv2.cuda.bilateralFilter(gpu_gray, 9, 75, 75)
                _, gpu_bin = cv2.cuda.threshold(gpu_gray, 180, 255, cv2.THRESH_BINARY)
                preproc_img = gpu_bin.download()
                img = Image.fromarray(preproc_img)
            page_txt = pytesseract.image_to_string(img, lang='fra+eng+ara')
            txt += f"\n--- PAGE {page_no} ---\n{page_txt}\n"
            blocks = re.findall(r"((?:[^\n]+\|[^\n]+\n)+)", page_txt)
            for block in blocks:
                rows = [x.split('|') for x in block.splitlines() if '|' in x]
                # Nettoie chaque cellule pour XLSX/XML
                rows = [[clean_xlsx_text(cell) for cell in row] for row in rows]
                if len(rows) > 1:
                    tables.append(rows)
        self.done.emit(txt, tables)

class ModernDarkOCRApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('🕶️ OCR PDF - PyQt5 (Darkly Modern)')
        self.resize(900, 850)
        self.layout = QVBoxLayout()
        self.setLayout(self.layout)

        dark_palette = QPalette()
        dark_palette.setColor(QPalette.Window, QColor(34, 34, 44))
        dark_palette.setColor(QPalette.Base, QColor(44, 44, 54))
        dark_palette.setColor(QPalette.Text, QColor(200, 200, 210))
        dark_palette.setColor(QPalette.Button, QColor(60, 60, 80))
        dark_palette.setColor(QPalette.ButtonText, QColor(232, 230, 255))
        dark_palette.setColor(QPalette.Highlight, QColor(80, 120, 255))
        dark_palette.setColor(QPalette.HighlightedText, QColor(40, 40, 40))
        self.setPalette(dark_palette)
        font_main = QFont("Segoe UI", 11)
        self.setFont(font_main)

        self.file_label = QLabel('Aucun fichier PDF sélectionné')
        self.file_label.setWordWrap(True)
        self.file_label.setStyleSheet("color: #FFD700; background: #2a2a2f; border-radius:7px; padding:8px; font-weight:bold;")
        self.layout.addWidget(self.file_label)

        self.pdf_btn = QPushButton('📄 Sélectionner PDF')
        self.pdf_btn.setStyleSheet("background:#314360; color:#AEEEEE; font-weight:bold; border-radius:18px; padding:9px;")
        self.pdf_btn.clicked.connect(self.select_pdf)
        self.layout.addWidget(self.pdf_btn)

        self.pages_edit = QLineEdit()
        self.pages_edit.setPlaceholderText("Pages à extraire (ex: 1,3-4,8 ou vide pour tout)")
        self.pages_edit.setStyleSheet("background:#2a2a2f; color:#FFD700; border-radius:8px; padding:7px;")
        self.layout.addWidget(self.pages_edit)

        self.batch_spin = QSpinBox()
        self.batch_spin.setMinimum(1)
        self.batch_spin.setMaximum(50)
        self.batch_spin.setValue(5)
        self.batch_spin.setStyleSheet("background:#2a2a2f; color:#FFD700; border-radius:8px; padding:6px;")
        self.layout.addWidget(QLabel('Taille de lot (pages traitées à la fois)'))
        self.layout.addWidget(self.batch_spin)

        self.progress_bar = QProgressBar()
        self.progress_bar.setMinimum(0)
        self.progress_bar.setMaximum(100)
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(True)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 2px solid #314360;
                border-radius: 8px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #1e90ff;
                width: 18px;
            }
        """)
        self.layout.addWidget(self.progress_bar)

        self.cancel_btn = QPushButton('⛔ Annuler')
        self.cancel_btn.setEnabled(False)
        self.cancel_btn.setStyleSheet("background:#c8004c; color:#fff; font-weight:bold; border-radius:15px; padding:6px;")
        self.cancel_btn.clicked.connect(self.cancel_ocr)
        self.layout.addWidget(self.cancel_btn)

        self.ocr_btn = QPushButton('🚦 Lancer OCR')
        self.ocr_btn.setStyleSheet("background:#28a745; color:#fff; font-weight:bold; border-radius:20px; padding:10px;")
        self.ocr_btn.clicked.connect(self.run_ocr)
        self.layout.addWidget(self.ocr_btn)

        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText("🔍 Recherche mot ou phrase")
        self.search_edit.setStyleSheet("background:#1a1a1f; color:#66D9EF; border-radius:8px; padding:6px;")
        self.layout.addWidget(self.search_edit)
        self.search_btn = QPushButton('Recherche')
        self.search_btn.setStyleSheet("background:#314360; color:#AEEEEE; font-weight:bold; border-radius:14px; padding:7px;")
        self.search_btn.clicked.connect(self.search_text)
        self.layout.addWidget(self.search_btn)

        self.result_box = QTextEdit()
        self.result_box.setAcceptRichText(False)
        self.result_box.setLineWrapMode(QTextEdit.WidgetWidth)
        self.result_box.setStyleSheet(
            "background:#181822; color:#C5C8FA; border-radius:10px; font-size:16px; padding:10px;")
        self.layout.addWidget(self.result_box)

        self.copy_btn = QPushButton('📋 Copier le texte')
        self.copy_btn.setStyleSheet("background:#314360; color:#FFD700; font-weight:bold; border-radius:16px; padding:8px;")
        self.copy_btn.clicked.connect(lambda: clipboard.copy(self.result_box.toPlainText()))
        self.layout.addWidget(self.copy_btn)

        self.export_docx_btn = QPushButton('💾 Exporter en DOCX')
        self.export_docx_btn.setStyleSheet("background:#1e90ff; color:#fff; font-weight:bold; border-radius:14px; padding:7px;")
        self.export_docx_btn.clicked.connect(self.export_docx)
        self.layout.addWidget(self.export_docx_btn)

        self.export_xlsx_btn = QPushButton('📊 Exporter tableaux en XLSX')
        self.export_xlsx_btn.setStyleSheet("background:#00bfff; color:#fff; font-weight:bold; border-radius:14px; padding:7px;")
        self.export_xlsx_btn.clicked.connect(self.export_xlsx)
        self.layout.addWidget(self.export_xlsx_btn)

        self.selected_file = ''
        self.tables = []
        self._pdf_num_pages = None

    def set_buttons_enabled(self, state):
        self.ocr_btn.setEnabled(state)
        self.pdf_btn.setEnabled(state)
        self.export_docx_btn.setEnabled(state)
        self.export_xlsx_btn.setEnabled(state)
        self.copy_btn.setEnabled(state)
        self.search_btn.setEnabled(state)
        self.cancel_btn.setEnabled(not state)

    def select_pdf(self):
        fname, _ = QFileDialog.getOpenFileName(self, "Choisir un PDF", "", "PDF Files (*.pdf)")
        if fname:
            self.selected_file = fname
            self.file_label.setText(fname)
            try:
                from PyPDF2 import PdfReader
                self._pdf_num_pages = len(PdfReader(open(fname, "rb")).pages)
            except:
                self._pdf_num_pages = None

    def run_ocr(self):
        if not self.selected_file:
            QMessageBox.warning(self, "Erreur", "Veuillez sélectionner un fichier PDF")
            return
        pages_str = self.pages_edit.text().strip()
        num_pages = self._pdf_num_pages
        if not num_pages:
            try:
                images = convert_from_path(self.selected_file, dpi=80)
                num_pages = len(images)
            except:
                QMessageBox.warning(self, "Erreur", "Impossible de lire le nombre de pages du PDF")
                return
        wanted_pages = parse_pages(pages_str, num_pages) if pages_str else list(range(1, num_pages + 1))
        self.result_box.clear()
        self.tables = []
        self._cur_pos = 0
        self._pages = wanted_pages
        self._batch_size = self.batch_spin.value()
        self._output = ""
        self._tables_output = []
        self.progress_bar.setValue(0)
        self.set_buttons_enabled(False)
        self.process_next_batch()

    def process_next_batch(self):
        start = self._cur_pos
        end = min(self._cur_pos + self._batch_size, len(self._pages))
        batch_pages = self._pages[start:end]
        if not batch_pages:
            self.result_box.setPlainText(self._output)
            self.tables = self._tables_output
            self.progress_bar.setValue(100)
            self.set_buttons_enabled(True)
            QMessageBox.information(self, "OCR terminé", f"OCR fini\nTableaux extraits : {len(self.tables)}")
            return
        total_pages = len(self._pages)
        self.worker = OCRWorker(self.selected_file, batch_pages, total_pages, self._cur_pos)
        self.worker.progress.connect(self.show_progress)
        self.worker.done.connect(self.on_batch_finish)
        self.worker.start()

    def show_progress(self, percent, msg):
        self.progress_bar.setValue(percent)
        self.file_label.setText(msg)

    def on_batch_finish(self, txt, tables):
        self._output += txt
        self._tables_output += tables
        self._cur_pos += len(self.worker.indices)
        self.result_box.setPlainText(self._output)
        
        old_worker = self.worker
        old_worker.deleteLater()
        self.worker = None

        if self._cur_pos < len(self._pages):
            self.process_next_batch()
        else:
            self.progress_bar.setValue(100)
            self.set_buttons_enabled(True)
            self.file_label.setText("✅ OCR terminé.")
            QMessageBox.information(self, "OCR terminé", f"OCR fini\nTableaux extraits : {len(self._tables_output)}")

    def cancel_ocr(self):
        if hasattr(self, 'worker') and self.worker and self.worker.isRunning():
            self.worker.cancel()
            self.worker.wait()
            self.file_label.setText("⛔ OCR annulé.")
            self.progress_bar.setValue(0)
            self.set_buttons_enabled(True)

    def search_text(self):
        keyword = self.search_edit.text()
        full_txt = self.result_box.toPlainText()
        count = len(re.findall(re.escape(keyword), full_txt, re.IGNORECASE))
        QMessageBox.information(self, "Résultat recherche", f"{count} occurrence(s) trouvée(s).")

    def export_docx(self):
        if not self.selected_file:
            return
        outname = self.selected_file + "_ocr.docx"
        doc = Document()
        safe_text = clean_docx_text(self.result_box.toPlainText())
        doc.add_paragraph(safe_text)
        doc.save(outname)
        QMessageBox.information(self, "Export DOCX", f"Sauvé : {outname}")

    def export_xlsx(self):
        if not self.selected_file or not self.tables:
            QMessageBox.warning(self, "Aucun tableau", "Aucun tableau détecté")
            return
        wb = openpyxl.Workbook()
        ws = wb.active
        for table in self.tables:
            for row in table:
                ws.append([clean_xlsx_text(cell) for cell in row])
            ws.append([])
        outname = self.selected_file + "_tables.xlsx"
        wb.save(outname)
        QMessageBox.information(self, "Export XLSX", f"Sauvé : {outname}")

    def closeEvent(self, event):
        if hasattr(self, 'worker') and self.worker.isRunning():
            self.worker.cancel()
            self.worker.wait()
        event.accept()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = ModernDarkOCRApp()
    window.show()
    sys.exit(app.exec_())
